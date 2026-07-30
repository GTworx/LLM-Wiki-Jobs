import os
import sys
import docx
import pypdf

sys.stdout.reconfigure(encoding='utf-8')

def extract_docx(file_path):
    doc = docx.Document(file_path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_txt:
                lines.append(row_txt)
    return "\n".join(lines)

def extract_pdf(file_path):
    reader = pypdf.PdfReader(file_path)
    lines = []
    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            lines.append(txt)
    return "\n".join(lines)

def main():
    resume_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "raw", "resume")
    files = os.listdir(resume_dir)
    for f in files:
        fp = os.path.join(resume_dir, f)
        print(f"Extracting {f}...")
        if f.endswith(".docx"):
            text = extract_docx(fp)
        elif f.endswith(".pdf"):
            text = extract_pdf(fp)
        else:
            continue
            
        out_name = f.replace(".docx", ".txt").replace(".pdf", ".txt")
        out_path = os.path.join(resume_dir, out_name)
        with open(out_path, "w", encoding="utf-8") as out:
            out.write(text)
        print(f"Saved extracted text to {out_path} ({len(text)} chars)")

if __name__ == "__main__":
    main()
